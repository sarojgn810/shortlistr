"""Extract resume text from uploaded files and convert to cv.md markdown."""

from __future__ import annotations

import io
import os
import re
import shutil
from typing import Any

from config import SHORTLISTR_ROOT, CV_MD_PATH

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md", ".markdown"}
MAX_UPLOAD_BYTES = 10 * 1024 * 1024  # 10 MB

_SECTION_ALIASES: dict[str, str] = {
    "professional summary": "PROFESSIONAL SUMMARY",
    "summary": "PROFESSIONAL SUMMARY",
    "profile": "PROFESSIONAL SUMMARY",
    "objective": "PROFESSIONAL SUMMARY",
    "core competencies": "CORE COMPETENCIES",
    "skills": "CORE COMPETENCIES",
    "technical skills": "CORE COMPETENCIES",
    "key skills": "CORE COMPETENCIES",
    "professional experience": "PROFESSIONAL EXPERIENCE",
    "work experience": "PROFESSIONAL EXPERIENCE",
    "experience": "PROFESSIONAL EXPERIENCE",
    "employment": "PROFESSIONAL EXPERIENCE",
    "education": "EDUCATION",
    "academic": "EDUCATION",
    "certifications": "CERTIFICATIONS",
    "certificates": "CERTIFICATIONS",
    "licenses": "CERTIFICATIONS",
    "projects": "PROJECTS",
    "achievements": "ACHIEVEMENTS",
}


def _ext(filename: str) -> str:
    return os.path.splitext(filename or "")[1].lower()


def validate_upload(filename: str, size: int) -> None:
    ext = _ext(filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext or '(none)'}'. Use PDF, DOCX, TXT, or Markdown."
        )
    if size <= 0:
        raise ValueError("Empty file.")
    if size > MAX_UPLOAD_BYTES:
        raise ValueError(f"File too large (max {MAX_UPLOAD_BYTES // (1024 * 1024)} MB).")


def _is_letter_spaced(line: str) -> bool:
    """True when a line arrived as "S e n i o r  D a t a" rather than words.

    Designer résumés often set headings and contact rows with letter-spacing,
    and PDF extractors emit each glyph's gap as a real space. Detected by the
    share of one-character tokens: normal prose has very few, a spaced line is
    almost all of them.
    """
    tokens = line.split()
    if len(tokens) < 6:
        return False
    singles = sum(1 for t in tokens if len(t) == 1)
    return singles / len(tokens) > 0.5


def _unspace_line(line: str) -> str:
    """Rebuild "J a n e  D o e" into "Jane Doe".

    Word boundaries survive extraction as a *wider* gap, so two or more spaces
    separate words and single spaces sit between the letters inside one. Any
    leading markdown marker is kept aside so "# a b c" does not become "#abc".
    """
    marker, body = "", line
    stripped = line.lstrip()
    if stripped[:1] in {"#", "-", "*", ">"}:
        head = len(line) - len(stripped)
        cut = len(stripped) - len(stripped.lstrip("#-*> "))
        marker, body = line[: head + cut], stripped[cut:]
    words = [re.sub(r"\s+", "", chunk) for chunk in re.split(r"\s{2,}", body)]
    return marker + " ".join(w for w in words if w)


_MONTHS = ("Jan", "Feb", "Mar", "Apr", "May", "Jun",
           "Jul", "Aug", "Sep", "Oct", "Nov", "Dec")


def normalize_spacing(text: str) -> str:
    """Put back spaces a design-heavy PDF never wrote.

    Typeset résumés place glyphs by position, so a contact row can carry no
    space characters at all and still look spaced on the page. pdfplumber
    reports what is actually in the file, which is why a real CV extracted as:

        Bangalore,India•+918884311573•realsarojnayak@gmail.com
        Feb2024–Present

    That is not a reader bug — the raw text is the same before and after the
    letter-spacing repair — but it reaches the rendered PDF, so it is worth
    correcting where the fix is unambiguous.

    Only clear cases are touched: separator bullets, a comma butted against a
    word, and a month butted against a year. Decimals, URLs and version numbers
    are left alone.
    """
    if not text:
        return text

    # Bullet and pipe separators used between contact fields. Horizontal space
    # only: \s* would swallow the newline before a bullet that starts a list
    # item and fold the whole list onto one line.
    text = re.sub(r"[ \t]*([•·|])[ \t]*", r" \1 ", text)
    # "Bangalore,India" — but not "1,200" or "a, b" which are already fine.
    text = re.sub(r"(?<=[A-Za-z]),(?=[A-Za-z])", ", ", text)
    # "Feb2024" -> "Feb 2024". Anchored to real month names so "H2024" and
    # "COVID19" are untouched.
    text = re.sub(rf"\b({'|'.join(_MONTHS)})(?=\d{{4}}\b)", r"\1 ", text)
    # Runs of spaces the substitutions may have doubled up.
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text


def rejoin_wrapped_lines(text: str) -> str:
    """Undo the hard wrapping a PDF page imposes on a flowing paragraph.

    A PDF has no paragraphs, only lines placed on a page, so a sentence that
    wrapped in the layout arrives with a newline in the middle of it:

        ... and Agentic AI —
        building LLM-powered autonomous incident response ...

    Rendered back out, that break appears mid-sentence for no reason.

    Only an unambiguous continuation is joined: the next line must begin with a
    lowercase letter, and the current one must not end a sentence. Headings,
    list items, table rows and blank lines are left exactly as they are, so
    document structure cannot be damaged by this.
    """
    lines = (text or "").split("\n")
    out: list[str] = []
    for line in lines:
        stripped = line.strip()
        prev = out[-1].strip() if out else ""
        continues = (
            out
            and prev
            and stripped
            and stripped[0].islower()
            and not prev.endswith((".", "!", "?", ":", ";"))
            and not prev.startswith(("#", "-", "*", ">", "|"))
            and not stripped.startswith(("#", "-", "*", ">", "|"))
        )
        if continues:
            out[-1] = out[-1].rstrip() + " " + stripped
        else:
            out.append(line)
    return "\n".join(out)


def repair_letter_spacing(text: str) -> str:
    """Undo per-glyph spacing, line by line, leaving normal lines untouched.

    Without this a letter-spaced résumé yields no email, no LinkedIn URL and no
    job title, because every pattern is looking for "@gmail.com" in a string
    that reads "@ g m a i l . c o m". On the résumé that surfaced this, 126 of
    136 lines were affected.
    """
    if not text:
        return text
    return "\n".join(
        _unspace_line(line) if _is_letter_spaced(line) else line
        for line in text.split("\n")
    )


def _pdf_text_pdfplumber(data: bytes) -> str:
    """Preferred reader: lays out glyphs by position, so letter-spacing survives.

    pypdf turns a letter-spaced résumé into "S e n i o r  D a t a  E n g i n e e r"
    and an email into "@ g m a i l . c o m", which no pattern can match. On the
    résumé that surfaced this, 126 of 136 lines came out that way.
    """
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text)
    return "\n".join(parts).strip()


def _pdf_text_pypdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n".join(parts).strip()


def extract_pdf_text(data: bytes) -> str:
    """Text from a PDF, best reader first.

    Letter-spacing repair still runs on whatever comes back: it is cheap, it
    leaves ordinary lines untouched, and the fallback reader needs it.
    """
    errors: list[str] = []
    for reader in (_pdf_text_pdfplumber, _pdf_text_pypdf):
        try:
            text = reader(data)
        except ImportError as e:
            errors.append(str(e))
            continue
        except Exception as e:  # a reader that chokes should not end the upload
            errors.append(f"{reader.__name__}: {e}")
            continue
        if text.strip():
            return rejoin_wrapped_lines(
                normalize_spacing(repair_letter_spacing(text))
            )

    if errors:
        raise RuntimeError(
            "Could not read this PDF — install a PDF reader with "
            "`pip install pdfplumber pypdf` (" + "; ".join(errors[:2]) + ")"
        )
    return ""


def extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("python-docx not installed — run: pip install python-docx") from e

    doc = Document(io.BytesIO(data))
    lines: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines).strip()


def extract_raw_text(filename: str, data: bytes) -> str:
    ext = _ext(filename)
    if ext == ".pdf":
        return extract_pdf_text(data)
    if ext in (".docx", ".doc"):
        return extract_docx_text(data)
    if ext in (".txt", ".md", ".markdown"):
        return data.decode("utf-8", errors="replace").strip()
    raise ValueError(f"Unsupported extension: {ext}")


def _is_section_heading(line: str) -> str | None:
    s = line.strip().rstrip(":").strip()
    if not s or len(s) > 80:
        return None
    key = re.sub(r"\s+", " ", s.lower())
    if key in _SECTION_ALIASES:
        return _SECTION_ALIASES[key]
    # ALL CAPS short line (common in PDF exports)
    letters = re.sub(r"[^A-Za-z]", "", s)
    if letters and letters.isupper() and 3 <= len(letters) <= 40 and len(s.split()) <= 6:
        return s.upper()
    return None


def _is_bullet_line(line: str) -> bool:
    return bool(re.match(r"^[\u2022\u2023\u25E6\u2043\u2219•\-\*●○◦]\s*", line)) or bool(
        re.match(r"^\d+[\.\)]\s+", line)
    )


def _strip_bullet(line: str) -> str:
    return re.sub(r"^[\u2022\u2023\u25E6\u2043\u2219•\-\*●○◦]\s*", "", line).strip()


def plain_text_to_markdown(text: str) -> str:
    """Heuristic PDF/DOCX plain text → structured markdown for cv.md."""
    if not text.strip():
        raise ValueError("No text could be extracted from this file.")

    # Already markdown-ish
    if re.search(r"^#{1,2}\s+\w", text, re.M):
        return text.strip()

    lines = [ln.rstrip() for ln in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    # Collapse excessive blank lines but keep paragraph breaks
    normalized: list[str] = []
    blank_run = 0
    for ln in lines:
        if not ln.strip():
            blank_run += 1
            if blank_run <= 1:
                normalized.append("")
            continue
        blank_run = 0
        normalized.append(ln.strip())

    out: list[str] = []
    i = 0
    # First substantive line → name
    while i < len(normalized) and not normalized[i].strip():
        i += 1
    if i < len(normalized):
        first = normalized[i].strip()
        # The first substantive line is the candidate's name. Treat it as the
        # title unless it is a *known* section heading — the generic ALL-CAPS
        # heuristic must not fire here, or an all-caps name (e.g. "ALEX CANDIDATE")
        # gets misread as a section instead of the H1 title.
        first_key = re.sub(r"\s+", " ", first.lower()).rstrip(":")
        if len(first) < 80 and first_key not in _SECTION_ALIASES:
            out.append(f"# {first}")
            i += 1
            # Next line often contact (email, phone, city)
            if i < len(normalized) and normalized[i].strip():
                contact = normalized[i].strip()
                if "@" in contact or re.search(r"\+?\d[\d\s\-()]{8,}", contact) or "|" in contact:
                    out.append(f"\n**{contact}**\n")
                    i += 1

    current_section: str | None = None
    while i < len(normalized):
        line = normalized[i]
        if not line:
            out.append("")
            i += 1
            continue

        heading = _is_section_heading(line)
        if heading:
            out.append(f"\n## {heading}\n")
            current_section = heading
            i += 1
            continue

        if _is_bullet_line(line):
            out.append(f"- {_strip_bullet(line)}")
        elif current_section == "PROFESSIONAL EXPERIENCE" and re.search(r"\d{4}", line):
            # Someone pasting markdown already has "### Company | Role | dates".
            # Prefixing again produced "### ### Company", and a downstream
            # lstrip("#") stops at the space — so the employer came out as
            # "### Infosys" and was shown to the candidate that way.
            out.append("\n" + (line if line.lstrip().startswith("#")
                               else f"### {line}"))
        else:
            out.append(line)
        i += 1

    md = "\n".join(out)
    md = re.sub(r"\n{3,}", "\n\n", md).strip()
    if not md.startswith("#"):
        name_line = ""
        for ln in normalized:
            if not ln.strip() or _is_section_heading(ln):
                continue
            if "@" in ln:
                continue
            if len(ln) < 80:
                name_line = ln
                break
        title = name_line or "Your Name"
        body = md if not name_line else re.sub(rf"^{re.escape(name_line)}\n?", "", md, count=1).strip()
        md = f"# {title}\n\n{body}".strip()
    return md


def save_resume_pdf(data: bytes, filename: str) -> str:
    """Copy uploaded PDF to repo root as resume.pdf (profile default)."""
    dest = os.path.join(SHORTLISTR_ROOT, "resume.pdf")
    with open(dest, "wb") as f:
        f.write(data)
    return dest


def ingest_resume_file(filename: str, data: bytes) -> dict[str, Any]:
    """Validate, extract, convert to markdown, optionally save PDF."""
    validate_upload(filename, len(data))
    raw = extract_raw_text(filename, data)
    markdown = plain_text_to_markdown(raw)
    pdf_path = None
    if _ext(filename) == ".pdf":
        pdf_path = save_resume_pdf(data, filename)

    with open(CV_MD_PATH, "w", encoding="utf-8") as f:
        f.write(markdown)

    return {
        "markdown": markdown,
        "source_format": _ext(filename).lstrip("."),
        "pdf_path": pdf_path,
        "cv_path": CV_MD_PATH,
        "char_count": len(markdown),
    }
